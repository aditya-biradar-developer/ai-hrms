"""
Resume Parser Service - Extract text from PDF/DOCX files
"""

import logging
import requests
import io
from PyPDF2 import PdfReader
from docx import Document
from typing import Optional

logger = logging.getLogger(__name__)

class ResumeParser:
    """Parse resumes from various file formats"""
    
    def parse_from_url(self, url: str) -> Optional[str]:
        """
        Fetch and parse resume from URL
        
        Args:
            url: URL to resume file (PDF or DOCX)
            
        Returns:
            Extracted text or None if failed
        """
        try:
            logger.info(f"📄 Fetching resume from: {url[:50]}...")
            
            # Convert Google Drive view links to direct download
            if 'drive.google.com' in url:
                url = self._convert_google_drive_url(url)
                logger.info(f"🔄 Converted to direct download: {url[:80]}...")
            
            # Download file with headers to avoid bot detection
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, timeout=30, allow_redirects=True, headers=headers)
            
            logger.info(f"📥 Response status: {response.status_code}")
            logger.info(f"📥 Content-Type: {response.headers.get('content-type', 'unknown')}")
            logger.info(f"📥 Content length: {len(response.content)} bytes")
            
            response.raise_for_status()
            
            # Check if we got HTML instead of a file (Google Drive error page)
            content_type = response.headers.get('content-type', '').lower()
            if 'text/html' in content_type and 'drive.google.com' in url:
                logger.error("❌ Got HTML instead of file - Google Drive access denied or file not public")
                logger.error(f"Response preview: {response.text[:200]}")
                return None
            
            # Detect file type from URL or content-type
            content_type = response.headers.get('content-type', '').lower()
            
            # Try to determine file type from URL first (more reliable than content-type)
            url_lower = url.lower()
            
            # Check URL or content-type for PDF
            if 'pdf' in url_lower or 'pdf' in content_type:
                logger.info("📄 Detected PDF file")
                return self._parse_pdf(response.content)
            
            # Check for DOCX
            elif 'docx' in url_lower or 'wordprocessingml' in content_type:
                logger.info("📄 Detected DOCX file")
                return self._parse_docx(response.content)
            
            # Check for old DOC format
            elif 'doc' in url_lower and 'docx' not in url_lower:
                logger.warning("⚠️ Old .doc format not fully supported, trying as DOCX")
                return self._parse_docx(response.content)
            
            # If content-type is generic (like octet-stream), try to guess from content
            elif 'octet-stream' in content_type or 'binary' in content_type:
                logger.info("📄 Generic content-type detected, trying to detect from content...")
                
                # Check magic bytes for PDF (starts with %PDF)
                if response.content[:4] == b'%PDF':
                    logger.info("✅ Detected PDF from magic bytes")
                    return self._parse_pdf(response.content)
                
                # Check magic bytes for ZIP (DOCX is a ZIP file)
                elif response.content[:2] == b'PK':
                    logger.info("✅ Detected ZIP/DOCX from magic bytes")
                    return self._parse_docx(response.content)
                
                else:
                    logger.error(f"❌ Could not detect file type from content")
                    logger.error(f"First 20 bytes: {response.content[:20]}")
                    return None
            
            else:
                logger.error(f"❌ Unsupported file type: {content_type}")
                return None
                
        except requests.RequestException as e:
            logger.error(f"❌ Failed to fetch resume: {e}")
            return None
        except Exception as e:
            logger.error(f"❌ Failed to parse resume: {e}")
            return None
    
    def _parse_pdf(self, content: bytes) -> Optional[str]:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                logger.warning("⚠️ PDF appears to be empty or image-based")
                return None
            
            logger.info(f"✅ Extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ PDF parsing error: {e}")
            return None
    
    def _parse_docx(self, content: bytes) -> Optional[str]:
        """Extract text from DOCX"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                logger.warning("⚠️ DOCX appears to be empty")
                return None
            
            logger.info(f"✅ Extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ DOCX parsing error: {e}")
            return None
    
    def _convert_google_drive_url(self, url: str) -> str:
        """Convert Google Drive view URL to direct download URL"""
        import re
        
        # Extract file ID from various Google Drive URL formats
        patterns = [
            r'/file/d/([a-zA-Z0-9_-]+)',
            r'id=([a-zA-Z0-9_-]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                file_id = match.group(1)
                return f'https://drive.google.com/uc?export=download&id={file_id}'
        
        return url  # Return original if no match
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        return '\n'.join(lines)
